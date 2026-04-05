import io
import logging
from datetime import datetime
from functools import wraps

from flask import Blueprint, jsonify, request, send_file, session

from models.settings import Setting

logger = logging.getLogger(__name__)
job_board_bp = Blueprint('job_board', __name__, url_prefix='/api/jobboard')


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return jsonify({'error': 'Unauthorized'}), 401
        return f(*args, **kwargs)
    return decorated


def _get_groq_client():
    import groq
    key = Setting.get('groq_api_key', '')
    if not key:
        raise ValueError('Groq API key not configured.')
    return groq.Groq(api_key=key)


# ── Public API ──────────────────────────────────────────────────────────────

@job_board_bp.get('/published')
def public_list():
    from utils.data_layer import jobpost_list
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 12))
    search = request.args.get('q', '').strip().lower()
    job_type = request.args.get('type', '').strip().lower()
    tag = request.args.get('tag', '').strip().lower()

    posts = jobpost_list(status='published')

    if search:
        posts = [p for p in posts if
                 search in (p.get('title') or '').lower()
                 or search in (p.get('company') or '').lower()
                 or search in (p.get('location') or '').lower()
                 or search in str(p.get('tags') or '').lower()
                 or search in (p.get('description') or '').lower()
                 or search in (p.get('original_description') or '').lower()]
    if job_type:
        posts = [p for p in posts if job_type in (p.get('job_type') or '').lower()]
    if tag:
        posts = [p for p in posts if tag in str(p.get('tags') or '').lower()]

    total = len(posts)
    start = (page - 1) * per_page
    page_posts = posts[start:start + per_page]
    return jsonify({
        'total': total,
        'page': page,
        'pages': max(1, -(-total // per_page)),
        'posts': page_posts,
    })


@job_board_bp.get('/published/<int:post_id>')
def public_detail(post_id):
    from utils.data_layer import jobpost_get
    post = jobpost_get(post_id)
    if not post or post.get('status') != 'published':
        return jsonify({'error': 'Not found'}), 404
    return jsonify(post)


def _ai_rewrite_job(post_dict: dict) -> str:
    from utils.ai_engine import rewrite_job_description
    src = (post_dict.get('original_description') or post_dict.get('description') or '').strip()
    if len(src) < 80:
        return src
    try:
        return rewrite_job_description(
            title=post_dict.get('title', ''),
            company=post_dict.get('company', ''),
            raw_description=src,
        )
    except Exception as e:
        logger.warning('AI rewrite skipped (%s – %s): %s', post_dict.get('title', '?'), post_dict.get('source', '?'), e)
        return src


@job_board_bp.post('/auto-rewrite')
@admin_required
def auto_rewrite():
    import concurrent.futures
    from utils.data_layer import jobpost_list_raw, jobpost_count, jobpost_update
    batch = min(int((request.get_json(silent=True) or {}).get('batch', 5)), 10)

    pending = jobpost_list_raw(status='published', ai_rewritten=False, limit=batch)
    total_remaining = jobpost_count(status='published', ai_rewritten=False)

    if not pending:
        return jsonify({'done': 0, 'remaining': 0, 'finished': True})

    def _call_ai(post_id, title, company, src):
        try:
            from utils.ai_engine import rewrite_job_description
            return post_id, rewrite_job_description(title=title, company=company, raw_description=src)
        except Exception as e:
            logger.warning('Auto-rewrite failed for job %s: %s', post_id, e)
            return post_id, None

    tasks = []
    for p in pending:
        src = (p.get('original_description') or p.get('description') or '').strip()
        if len(src) >= 80:
            tasks.append((p['id'], p.get('title') or '', p.get('company') or '', src))

    rewrite_map = {}
    if tasks:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(tasks), 5)) as ex:
            futures = {ex.submit(_call_ai, pid, t, c, src): pid for pid, t, c, src in tasks}
            for f in concurrent.futures.as_completed(futures, timeout=60):
                try:
                    pid, rewritten = f.result()
                    rewrite_map[pid] = rewritten
                except Exception as e:
                    logger.warning('Future failed: %s', e)

    done = 0
    for post in pending:
        pid = post['id']
        update_data = {'ai_rewritten': True}
        if pid in rewrite_map and rewrite_map[pid]:
            update_data['description'] = rewrite_map[pid]
        jobpost_update(pid, update_data)
        done += 1

    still_remaining = jobpost_count(status='published', ai_rewritten=False)
    return jsonify({'done': done, 'remaining': still_remaining, 'finished': still_remaining == 0})


@job_board_bp.get('/rewrite-status')
@admin_required
def rewrite_status():
    from utils.data_layer import jobpost_count
    pending = jobpost_count(status='published', ai_rewritten=False)
    total = jobpost_count(status='published')
    return jsonify({'pending': pending, 'total': total, 'finished': pending == 0})


@job_board_bp.get('/live-search')
def live_search():
    import concurrent.futures
    from utils.data_layer import jobpost_list
    from utils.job_aggregator import fetch_remotive, fetch_arbeitnow, fetch_remoteok

    q = request.args.get('q', '').strip()
    job_type = request.args.get('type', '').strip().lower()
    search = q.lower()

    try:
        db_posts = jobpost_list(status='published')
    except Exception as e:
        logger.warning('Could not load local posts for live-search: %s', e)
        db_posts = []
    if search:
        db_posts = [p for p in db_posts if
                    search in (p.get('title') or '').lower()
                    or search in (p.get('company') or '').lower()
                    or search in (p.get('location') or '').lower()
                    or search in str(p.get('tags') or '').lower()
                    or search in (p.get('description') or '').lower()
                    or search in (p.get('original_description') or '').lower()]
    if job_type:
        db_posts = [p for p in db_posts if job_type in (p.get('job_type') or '').lower()]

    local_results = []
    for p in db_posts:
        d = dict(p)
        d['is_live'] = False
        local_results.append(d)

    live_results = []
    if q:
        def _remotive():
            try:
                return fetch_remotive(search=q, limit=10)
            except Exception:
                return []

        def _arbeitnow():
            try:
                raw = fetch_arbeitnow(limit=30)
                return [r for r in raw if search in (r.get('title') or '').lower()
                        or search in (r.get('company') or '').lower()
                        or search in str(r.get('tags') or '').lower()
                        or search in (r.get('original_description') or '').lower()][:10]
            except Exception:
                return []

        def _remoteok():
            try:
                raw = fetch_remoteok(limit=20)
                return [r for r in raw if search in (r.get('title') or '').lower()
                        or search in (r.get('company') or '').lower()
                        or search in str(r.get('tags') or '').lower()
                        or search in (r.get('original_description') or '').lower()][:10]
            except Exception:
                return []

        raw_live = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
            fetch_futures = [ex.submit(_remotive), ex.submit(_arbeitnow), ex.submit(_remoteok)]
            for f in concurrent.futures.as_completed(fetch_futures, timeout=10):
                try:
                    raw_live.extend(f.result())
                except Exception:
                    pass

        local_ext_ids = {p.get('external_id') for p in db_posts if p.get('external_id')}
        local_keys = {(p.get('title') or '').lower() + '|' + (p.get('company') or '').lower() for p in db_posts}
        seen_live = set()
        deduped_live = []
        for r in raw_live:
            ext_id = r.get('external_id', '')
            key = (r.get('title') or '').lower() + '|' + (r.get('company') or '').lower()
            if ext_id and ext_id in local_ext_ids:
                continue
            if key in local_keys or key in seen_live:
                continue
            seen_live.add(key)
            r['is_live'] = True
            r['id'] = None
            r['description'] = r.get('original_description') or r.get('description') or ''
            deduped_live.append(r)

        if job_type:
            deduped_live = [r for r in deduped_live if job_type in (r.get('job_type') or '').lower()]

        live_results = deduped_live

    all_results = local_results + live_results
    return jsonify({
        'total': len(all_results),
        'local_count': len(local_results),
        'live_count': len(live_results),
        'results': all_results,
    })


@job_board_bp.get('/featured')
def featured_posts():
    from utils.data_layer import jobpost_list
    posts = jobpost_list(status='published', featured=True, limit=6)
    if not posts:
        posts = jobpost_list(status='published', limit=6)
    return jsonify(posts)


# ── Export API ──────────────────────────────────────────────────────────────

@job_board_bp.post('/export')
@admin_required
def export_jobs():
    from utils.data_layer import jobpost_list
    data = request.get_json(silent=True) or {}
    ids = data.get('ids', [])
    fmt = data.get('format', 'txt').lower()
    search = data.get('search', '').strip().lower()
    job_type = data.get('job_type', '').strip().lower()

    posts = jobpost_list(status='published')

    if ids:
        id_set = {int(i) for i in ids if str(i).isdigit()}
        posts = [p for p in posts if p.get('id') in id_set]
    else:
        if search:
            posts = [p for p in posts if
                     search in (p.get('title') or '').lower()
                     or search in (p.get('company') or '').lower()
                     or search in (p.get('location') or '').lower()
                     or search in str(p.get('tags') or '').lower()
                     or search in (p.get('description') or '').lower()
                     or search in (p.get('original_description') or '').lower()]
        if job_type:
            posts = [p for p in posts if job_type in (p.get('job_type') or '').lower()]

    if not posts:
        return jsonify({'error': 'No jobs to export'}), 400

    def _tags_str(p):
        tags = p.get('tags', [])
        if isinstance(tags, list):
            return ', '.join(tags)
        return tags or ''

    if fmt == 'txt':
        lines = []
        for p in posts:
            lines.append('=' * 60)
            lines.append(f'Title:    {p.get("title", "")}')
            lines.append(f'Company:  {p.get("company") or "N/A"}')
            lines.append(f'Location: {p.get("location") or "N/A"}')
            lines.append(f'Type:     {p.get("job_type") or "N/A"}')
            if p.get('salary'):
                lines.append(f'Salary:   {p["salary"]}')
            if p.get('apply_url'):
                lines.append(f'Apply:    {p["apply_url"]}')
            tags = _tags_str(p)
            if tags:
                lines.append(f'Tags:     {tags}')
            desc = (p.get('description') or p.get('original_description') or '').strip()
            if desc:
                lines.append('')
                lines.append(desc[:800])
            lines.append('')
        content = '\n'.join(lines).encode('utf-8')
        buf = io.BytesIO(content)
        buf.seek(0)
        return send_file(buf, mimetype='text/plain',
                         as_attachment=True, download_name='job_listings.txt')

    elif fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading('Job Listings Export', 0)
        for p in posts:
            h = doc.add_heading(p.get('title', ''), level=1)
            h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            details = []
            if p.get('company'):  details.append(f'Company: {p["company"]}')
            if p.get('location'): details.append(f'Location: {p["location"]}')
            if p.get('job_type'): details.append(f'Type: {p["job_type"]}')
            if p.get('salary'):   details.append(f'Salary: {p["salary"]}')
            if p.get('apply_url'): details.append(f'Apply: {p["apply_url"]}')
            tags = _tags_str(p)
            if tags: details.append(f'Tags: {tags}')
            for d in details:
                para = doc.add_paragraph(d)
                para.runs[0].bold = True
            desc = (p.get('description') or p.get('original_description') or '').strip()
            if desc:
                doc.add_paragraph(desc[:800])
            doc.add_paragraph('')
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name='job_listings.docx')

    elif fmt == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        buf = io.BytesIO()
        doc_pdf = SimpleDocTemplate(buf, pagesize=A4,
                                    leftMargin=20*mm, rightMargin=20*mm,
                                    topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('jb_title', parent=styles['Heading1'],
                                     textColor=colors.HexColor('#4F46E5'), fontSize=14, spaceAfter=4)
        meta_style = ParagraphStyle('jb_meta', parent=styles['Normal'],
                                    textColor=colors.HexColor('#475569'), fontSize=9, spaceAfter=2)
        desc_style = ParagraphStyle('jb_desc', parent=styles['Normal'],
                                    fontSize=9, leading=13, spaceAfter=6,
                                    textColor=colors.HexColor('#334155'))
        story = [Paragraph('Job Listings Export', styles['Title']), Spacer(1, 8*mm)]
        for p in posts:
            story.append(Paragraph(p.get('title') or 'Untitled', title_style))
            meta_parts = []
            if p.get('company'):   meta_parts.append(f'<b>Company:</b> {p["company"]}')
            if p.get('location'):  meta_parts.append(f'<b>Location:</b> {p["location"]}')
            if p.get('job_type'):  meta_parts.append(f'<b>Type:</b> {p["job_type"]}')
            if p.get('salary'):    meta_parts.append(f'<b>Salary:</b> {p["salary"]}')
            if p.get('apply_url'): meta_parts.append(f'<b>Apply:</b> {p["apply_url"]}')
            tags = _tags_str(p)
            if tags: meta_parts.append(f'<b>Tags:</b> {tags}')
            for m in meta_parts:
                story.append(Paragraph(m, meta_style))
            desc = (p.get('description') or p.get('original_description') or '').strip()
            if desc:
                safe_desc = desc[:800].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_desc, desc_style))
            story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#E2E8F0')))
            story.append(Spacer(1, 5*mm))
        doc_pdf.build(story)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name='job_listings.pdf')

    return jsonify({'error': 'Invalid format'}), 400


# ── Admin API ───────────────────────────────────────────────────────────────

@job_board_bp.get('/admin/posts')
@admin_required
def admin_list():
    from utils.data_layer import jobpost_list, jobpost_count_by_status
    status = request.args.get('status', 'all')
    if status == 'all':
        posts = jobpost_list()
    else:
        posts = jobpost_list(status=status)
    counts = jobpost_count_by_status()
    return jsonify({'posts': posts, 'counts': counts})


@job_board_bp.get('/admin/posts/<int:post_id>')
@admin_required
def admin_get(post_id):
    from utils.data_layer import jobpost_get
    post = jobpost_get(post_id)
    if post is None:
        return jsonify({'error': 'Not found'}), 404
    post['original_description'] = post.get('original_description', '')
    post['description_raw'] = post.get('description', '')
    return jsonify(post)


@job_board_bp.put('/admin/posts/<int:post_id>')
@admin_required
def admin_update(post_id):
    from utils.data_layer import jobpost_get, jobpost_update
    post = jobpost_get(post_id)
    if post is None:
        return jsonify({'error': 'Not found'}), 404
    data = request.get_json(silent=True) or {}
    update_data = {}
    for field in ['title', 'company', 'location', 'job_type', 'salary', 'apply_url', 'description']:
        if field in data:
            update_data[field] = data[field]
    if 'tags' in data:
        tags = data['tags']
        update_data['tags'] = ', '.join(tags) if isinstance(tags, list) else tags
    if 'featured' in data:
        update_data['featured'] = bool(data['featured'])
    if 'status' in data and data['status'] in ('draft', 'published', 'archived'):
        update_data['status'] = data['status']
    updated = jobpost_update(post_id, update_data)
    return jsonify({'success': True, 'post': updated})


@job_board_bp.delete('/admin/posts/<int:post_id>')
@admin_required
def admin_delete(post_id):
    from utils.data_layer import jobpost_delete
    if not jobpost_delete(post_id):
        return jsonify({'error': 'Not found'}), 404
    return jsonify({'success': True})


@job_board_bp.post('/admin/posts/bulk')
@admin_required
def admin_bulk():
    from utils.data_layer import jobpost_bulk
    data = request.get_json(silent=True) or {}
    action = data.get('action')
    ids = data.get('ids', [])
    if not ids or action not in ('publish', 'archive', 'draft', 'delete'):
        return jsonify({'success': False, 'error': 'Invalid request'}), 400
    affected = jobpost_bulk(ids, action)
    return jsonify({'success': True, 'affected': affected})


@job_board_bp.post('/admin/posts/<int:post_id>/rewrite')
@admin_required
def admin_rewrite(post_id):
    from utils.data_layer import jobpost_get, jobpost_update
    post = jobpost_get(post_id)
    if post is None:
        return jsonify({'error': 'Not found'}), 404
    source_text = post.get('original_description') or post.get('description')
    if not source_text or len(source_text.strip()) < 50:
        return jsonify({'success': False, 'error': 'Job description is too short to rewrite.'}), 400
    try:
        from utils.ai_engine import rewrite_job_description
        rewritten = rewrite_job_description(
            title=post.get('title') or '',
            company=post.get('company') or '',
            raw_description=source_text,
        )
        jobpost_update(post_id, {'description': rewritten, 'ai_rewritten': True})
        return jsonify({'success': True, 'description': rewritten})
    except Exception as e:
        logger.error('Groq rewrite error: %s', e)
        return jsonify({'success': False, 'error': str(e)}), 500


@job_board_bp.post('/admin/fetch')
@admin_required
def admin_fetch():
    from utils.data_layer import jobpost_find_by_external_id, jobpost_create, jobpost_count
    data = request.get_json(silent=True) or {}
    sources = data.get('sources', ['remotive', 'arbeitnow', 'remoteok'])
    search = data.get('search', '')
    limit = int(data.get('limit_per_source', 15))

    adzuna_id = Setting.get('adzuna_app_id', '')
    adzuna_key = Setting.get('adzuna_app_key', '')

    try:
        from utils.job_aggregator import aggregate
        posts_data = aggregate(
            sources=sources,
            search=search,
            limit_per_source=limit,
            adzuna_app_id=adzuna_id or None,
            adzuna_app_key=adzuna_key or None,
        )
    except Exception as e:
        logger.error('Aggregation error: %s', e)
        return jsonify({'success': False, 'error': str(e)}), 500

    added = 0
    skipped = 0
    for p in posts_data:
        ext_id = p.get('external_id', '')
        if ext_id and jobpost_find_by_external_id(ext_id):
            skipped += 1
            continue
        desc = p.get('original_description', '')
        jobpost_create({
            'external_id': ext_id,
            'source': p.get('source', 'unknown'),
            'title': p.get('title', ''),
            'company': p.get('company', ''),
            'location': p.get('location', ''),
            'job_type': p.get('job_type', ''),
            'salary': p.get('salary', ''),
            'tags': p.get('tags', ''),
            'apply_url': p.get('apply_url', ''),
            'original_description': desc,
            'description': desc,
            'status': 'draft',
        })
        added += 1

    pending_rewrite = jobpost_count(status='published', ai_rewritten=False)
    return jsonify({
        'success': True,
        'added': added,
        'skipped': skipped,
        'total': len(posts_data),
        'pending_rewrite': pending_rewrite,
    })
